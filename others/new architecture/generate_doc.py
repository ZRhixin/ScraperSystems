"""
Generates the new architecture plan as a .docx file.
Run once: python generate_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = val
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].bold = True
    return row

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Chain-of-Title Scraper\nNew Architecture Plan')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('The Research Team Pattern — Multi-Agent Hierarchy').font.size = Pt(14)

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run('May 2026').font.size = Pt(11)

doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. Executive Summary
# ---------------------------------------------------------------------------
h1('1. Executive Summary')

body(
    'The current architecture uses a linear pipeline — Scout → Investigate → Conclude → Verify — '
    'where a single Investigator agent is responsible for querying county assessors, searching '
    'Register of Deeds, pulling PDFs, reading documents, reasoning about chain of title, '
    'managing court records, and writing to the database. This creates an overloaded agent that '
    'stops early, misses phases, and is expensive to run.'
)

body(
    'The new architecture replaces this with a Research Team pattern: one Case Manager orchestrator '
    'that delegates to specialized sub-agents, each an expert in exactly one domain. The business '
    'logic, database schema, and prompts are preserved. Only the organization of intelligence changes.'
)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 2. Problem With Current Architecture
# ---------------------------------------------------------------------------
h1('2. Problems With the Current Architecture')

problems = [
    ('Overloaded Investigator',
     'One agent has 15 tools and a 4,000-word system prompt covering deeds, courts, OCR, '
     'chain reasoning, and database writes. It consistently stops at Phase A and gives up.'),
    ('Timeout Risk',
     'The entire Scout → Investigate → Conclude → Verify chain runs synchronously on one '
     'webhook. Investigate alone takes 10–18 minutes. HTTP callers time out before completion.'),
    ('No Loopback',
     'When Verify raises objections, nothing routes back to Investigate. The feedback loop '
     'described in the build spec is not implemented.'),
    ('Wrong Model for Each Task',
     'GPT-5 runs on Scout (pure extraction) and Investigate (multi-step loop). '
     'Expensive models run on every step regardless of complexity.'),
    ('County Scalability',
     'Adding a new county requires modifying the Investigator system prompt and rewiring tools. '
     'There is no clean separation between county-specific retrieval and chain reasoning.'),
    ('No Parallelism',
     'Every step runs sequentially. Deeds searches and court searches that could run in parallel '
     'are forced into a linear queue.'),
]

for title_text, desc in problems:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(f'{title_text}: ')
    run_bold.bold = True
    p.add_run(desc)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 3. New Architecture Overview
# ---------------------------------------------------------------------------
h1('3. New Architecture — The Research Team')

body(
    'The new architecture models how a real title company operates. Senior experts delegate '
    'retrieval work to specialists. The reasoning layer only sees clean summaries, not raw '
    'HTML and county website responses.'
)

doc.add_paragraph()
h2('3.1 Agent Roster')

agents = [
    ('Case Manager', 'Orchestrator', 'n8n routing logic (no AI)',
     'Receives parcel_id, delegates to sub-agents in sequence, enforces loop limits, tracks job status.'),
    ('Property Researcher', 'Scout replacement', 'GPT-4o-mini',
     'Queries county assessors only. Extracts owner, legal description, transfer history. Writes to DB.'),
    ('Deeds Expert', 'Retrieval specialist', 'GPT-4o-mini',
     'Searches Register of Deeds by name or book/page. Pulls deed PDFs. Saves captures. No chain reasoning.'),
    ('Court Expert', 'Retrieval specialist', 'GPT-4o-mini',
     'Searches NC Clerk of Superior Court. Pulls estate files, foreclosure records, judgments. Saves captures.'),
    ('Document Analyst', 'Extraction specialist', 'Claude Haiku',
     'Reads any capture. Runs OCR extraction. Returns structured fields and flags. No retrieval.'),
    ('Title Attorney', 'Chain Analyst (Investigator replacement)', 'Claude Sonnet',
     'The only reasoning-heavy agent. Directs specialists, reasons about chain of title, settles or flags the case.'),
    ('Conclusion Writer', 'Conclude replacement', 'Claude Haiku',
     'Reads DB state after investigation. Produces structured chain_conclusions JSON.'),
    ('Senior Partner', 'Verifier replacement', 'Claude Haiku',
     'Adversarial review of conclusion against evidence only. No investigation trace access.'),
]

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
header = tbl.rows[0]
header.cells[0].text = 'Agent'
header.cells[1].text = 'Role'
header.cells[2].text = 'Model'
header.cells[3].text = 'Responsibility'
for cell in header.cells:
    cell.paragraphs[0].runs[0].bold = True
shade_row(header, 'D9E1F2')

for agent_data in agents:
    row = tbl.add_row()
    for i, val in enumerate(agent_data):
        row.cells[i].text = val

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 4. Agent Detail
# ---------------------------------------------------------------------------
h1('4. Agent Detail')

# --- Property Researcher ---
h2('4.1 Property Researcher')
bullet('Single responsibility: query county assessors')
bullet('Tools: Wake Assessor, Mecklenburg Assessor, Buncombe Assessor, New Hanover Assessor')
bullet('Input: parcel_id, county')
bullet('Output: structured property data written to DB, returns property_id')
bullet('Model: GPT-4o-mini — pure extraction, zero chain reasoning required')
bullet('Adding a new county: add one tool node. No other agent changes.')

doc.add_paragraph()

# --- Deeds Expert ---
h2('4.2 Deeds Expert')
bullet('Single responsibility: find and retrieve deed records from Register of Deeds')
bullet('Tools: Wake Deeds Search, Pull Deed (all counties), Save Capture')
bullet('Input: task description — search by name, or pull by book + page')
bullet('Output: list of captures with grantor, grantee, doc type, book, page')
bullet('Does NOT reason about chain of title — only retrieves documents')
bullet('Model: GPT-4o-mini — search and retrieval only')
body('')
body(
    'The Deeds Expert handles all county-specific mechanics. The Title Attorney never '
    'needs to know whether it is searching Wake or Mecklenburg — it just asks the '
    'Deeds Expert for what it needs.'
)

doc.add_paragraph()

# --- Court Expert ---
h2('4.3 Court Expert')
bullet('Single responsibility: search and retrieve NC court records')
bullet('Tools: Court Search, Register of Actions, Court Pull')
bullet('Input: owner name, county, context about what to look for')
bullet('Output: estate files, foreclosure stages, judgments, liens')
bullet('Does NOT reason about ownership implications — retrieves and summarizes only')
bullet('Model: GPT-4o-mini — retrieval only')

doc.add_paragraph()

# --- Document Analyst ---
h2('4.4 Document Analyst')
bullet('Single responsibility: read a single capture and extract structured data')
bullet('Tools: Read Document (OCR + extraction subroutine)')
bullet('Input: capture_id, parcel reference')
bullet('Output: structured extraction — grantor, grantee, vesting language, legal description, flags')
bullet('Does not search for documents — only reads what it is given')
bullet('Model: Claude Haiku — structured extraction from OCR text')

doc.add_paragraph()

# --- Title Attorney ---
h2('4.5 Title Attorney (Chain Analyst)')
body(
    'This is the only reasoning-heavy agent. It replaces the current Investigator '
    'but has a fundamentally lighter tool set — it delegates all retrieval to specialists '
    'and focuses entirely on chain logic.'
)
doc.add_paragraph()
bullet('Tools: ask_deeds_expert, ask_court_expert, ask_document_analyst, log_trace, '
       'open_question, resolve_question, settle_chain, flag_review')
bullet('Input: property_id, full property state from DB, optional objection context from Verifier')
bullet('Does NOT directly call county websites — delegates to Deeds Expert and Court Expert')
bullet('Focuses on: which deeds to request, how they connect, what gaps exist, when to settle')
bullet('Model: Claude Sonnet — the only place where expensive reasoning is justified')
doc.add_paragraph()
body(
    'Because the Title Attorney only sees structured summaries from sub-agents '
    '(not raw HTML or county website responses), its context window stays clean. '
    'It can reason about a 5-deed chain without being cluttered by raw OCR output.'
)

doc.add_paragraph()

# --- Conclusion Writer ---
h2('4.6 Conclusion Writer')
bullet('Single responsibility: produce the chain_conclusions JSON from settled investigation data')
bullet('Reads DB: property row, settled session, all extractions, incidentals, aggregated flags')
bullet('No tools — DB read is handled by the /conclude/data endpoint before calling this agent')
bullet('Model: Claude Haiku — single-shot structured output, no reasoning loops')

doc.add_paragraph()

# --- Senior Partner ---
h2('4.7 Senior Partner (Verifier)')
bullet('Single responsibility: adversarial review of the conclusion against document evidence')
bullet('Input: chain_conclusions row + all referenced document_extractions')
bullet('Does NOT see investigation_trace — evaluates evidence only')
bullet('Produces verdict: approved / objection_raised / flagged_for_human')
bullet('Model: Claude Haiku — structured adversarial check')

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 5. Case Manager Flow
# ---------------------------------------------------------------------------
h1('5. Case Manager — Orchestration Flow')

body(
    'The Case Manager is not an AI agent. It is n8n routing logic that sequences '
    'the sub-agents and enforces business rules.'
)

doc.add_paragraph()

steps = [
    '1. Receive parcel_id + county via webhook',
    '2. Respond immediately with { job_id } — caller does not wait',
    '3. Call Property Researcher(parcel_id, county) → get property_id',
    '4. Call Title Attorney(property_id) → runs until settled or flagged',
    '5. Call Conclusion Writer(property_id) → get conclusion_id',
    '6. Call Verifier(conclusion_id) → get verdict',
    '',
    '   IF verdict = "objection_raised" AND loop_count < 1:',
    '     Call Title Attorney(property_id, objections) — passes objection context',
    '     Call Conclusion Writer(property_id) → new conclusion_id',
    '     Call Verifier(new conclusion_id) → new verdict',
    '',
    '   IF verdict still not "approved" OR loop_count >= 1:',
    '     Force status = flagged_for_human',
    '',
    '7. Write final status to DB',
    '8. Client polls GET /job/{job_id} for status and result',
]

for step in steps:
    if step == '':
        doc.add_paragraph()
    elif step.startswith('   '):
        p = doc.add_paragraph(step.strip(), style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.75)
    else:
        doc.add_paragraph(step, style='List Bullet')

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 6. Tool Hierarchy
# ---------------------------------------------------------------------------
h1('6. Tool Hierarchy')

body(
    'Each agent only has tools appropriate to its role. No agent can reach outside its domain.'
)

doc.add_paragraph()

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
h = tbl2.rows[0]
h.cells[0].text = 'Agent'
h.cells[1].text = 'Can Call'
h.cells[2].text = 'Cannot Call'
for cell in h.cells:
    cell.paragraphs[0].runs[0].bold = True
shade_row(h, 'D9E1F2')

tool_rows = [
    ('Case Manager', 'Property Researcher, Title Attorney, Conclusion Writer, Verifier',
     'County websites, DB directly'),
    ('Property Researcher', 'Wake/Meck/Buncombe/NewHanover Assessor endpoints',
     'Deeds, Court, Document Reader'),
    ('Deeds Expert', 'Wake Deeds Search, Pull Deed, Save Capture',
     'Assessor, Court, Document Reader'),
    ('Court Expert', 'Court Search, Register of Actions, Court Pull',
     'Assessor, Deeds, Document Reader'),
    ('Document Analyst', 'Read Document (OCR + extraction)',
     'Assessor, Deeds, Court'),
    ('Title Attorney', 'ask_deeds_expert, ask_court_expert, ask_document_analyst, '
                       'log_trace, open/resolve_question, settle_chain, flag_review',
     'County websites directly'),
    ('Conclusion Writer', '/conclude/data endpoint (read only)',
     'All retrieval tools'),
    ('Senior Partner', '/verify/data endpoint (read only)',
     'All retrieval tools, investigation_trace'),
]

for row_data in tool_rows:
    row = tbl2.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 7. Model Cost Comparison
# ---------------------------------------------------------------------------
h1('7. Model Cost Comparison')

tbl3 = doc.add_table(rows=1, cols=4)
tbl3.style = 'Table Grid'
h3_row = tbl3.rows[0]
for i, txt in enumerate(['Agent', 'Current Model', 'New Model', 'Reason']):
    h3_row.cells[i].text = txt
    h3_row.cells[i].paragraphs[0].runs[0].bold = True
shade_row(h3_row, 'D9E1F2')

cost_rows = [
    ('Property Researcher', 'GPT-5 (Scout)', 'GPT-4o-mini', 'Pure extraction, no reasoning'),
    ('Deeds Expert', 'GPT-5 (Investigator)', 'GPT-4o-mini', 'Search and retrieval only'),
    ('Court Expert', 'GPT-5 (Investigator)', 'GPT-4o-mini', 'Search and retrieval only'),
    ('Document Analyst', 'Claude Sonnet 4.6', 'Claude Haiku 4.5', 'Structured extraction, single shot'),
    ('Title Attorney', 'GPT-5 (Investigator)', 'Claude Sonnet 4.6', 'Chain reasoning — expensive model justified here only'),
    ('Conclusion Writer', 'Claude Opus 4.7', 'Claude Haiku 4.5', 'Single-shot structured JSON output'),
    ('Senior Partner', 'Claude Opus 4.7', 'Claude Haiku 4.5', 'Single-shot adversarial check'),
]

for row_data in cost_rows:
    row = tbl3.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_paragraph()

body(
    'The key shift: only the Title Attorney uses an expensive model. Every retrieval task '
    'and every single-shot output uses a cheap model. The current architecture runs GPT-5 '
    'through up to 50 agent iterations on the Investigator, regardless of what it is doing.'
)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 8. Parallelism Opportunities
# ---------------------------------------------------------------------------
h1('8. Parallelism Opportunities')

body('The Case Manager can dispatch sub-agents in parallel where tasks are independent:')
doc.add_paragraph()

bullet('On intake: Property Researcher + initial Deeds name search can run in parallel '
       'once parcel_id is known.')
bullet('During Title Attorney execution: Deeds Expert and Court Expert can run searches '
       'simultaneously when both are needed (e.g., Phase B + Phase D together).')
bullet('Document Analyst can process multiple captures in parallel rather than one at a time.')

doc.add_paragraph()
body(
    'The current architecture is fully sequential. Every deed pull waits for the previous '
    'one to complete. With specialist agents, the Title Attorney can dispatch 3 deed pulls '
    'simultaneously and process them all when results return.'
)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 9. County Scalability
# ---------------------------------------------------------------------------
h1('9. County Scalability')

body('Adding a new county under the new architecture:')
doc.add_paragraph()

bullet('Add one HTTP tool to Property Researcher pointing to the new assessor endpoint')
bullet('Add the county to Deeds Expert\'s tool list (or update its system prompt)')
bullet('Title Attorney does not change')
bullet('Case Manager does not change')
bullet('Conclusion Writer does not change')
bullet('Verifier does not change')

doc.add_paragraph()

body(
    'Under the current architecture, adding a county requires modifying the Investigator '
    'system prompt, adding tools to the Investigator node, and verifying the agent still '
    'behaves correctly with the expanded tool set. The Investigator already has 15 tools — '
    'each addition increases the risk of tool selection errors.'
)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 10. Database Changes
# ---------------------------------------------------------------------------
h1('10. Database Changes')

body('No schema changes are required. The existing tables cover all needs:')
doc.add_paragraph()

bullet('properties, appraiser_transfer_history — written by Property Researcher')
bullet('rod_captures, court_captures — written by Deeds Expert and Court Expert')
bullet('document_extractions — written by Document Analyst')
bullet('investigation_sessions, investigation_trace, investigation_questions — written by Title Attorney')
bullet('incidental_records — written by Title Attorney via log_incidental')
bullet('chain_conclusions — written by Conclusion Writer')

doc.add_paragraph()

body(
    'One addition recommended: a jobs table for async status tracking.'
)

doc.add_paragraph()

tbl4 = doc.add_table(rows=1, cols=2)
tbl4.style = 'Table Grid'
hj = tbl4.rows[0]
hj.cells[0].text = 'Column'
hj.cells[1].text = 'Description'
for cell in hj.cells:
    cell.paragraphs[0].runs[0].bold = True
shade_row(hj, 'D9E1F2')

job_cols = [
    ('id', 'UUID — returned to caller on job submission'),
    ('property_id', 'FK to properties — populated after Property Researcher completes'),
    ('status', 'queued | running | conclude | verify | complete | failed | flagged_for_human'),
    ('current_agent', 'Which sub-agent is currently executing'),
    ('loop_count', 'How many Investigate → Conclude → Verify loops have run'),
    ('result', 'JSONB — final output when complete'),
    ('created_at, updated_at', 'Standard timestamps'),
]

for col_data in job_cols:
    row = tbl4.add_row()
    row.cells[0].text = col_data[0]
    row.cells[1].text = col_data[1]

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 11. Trade-offs
# ---------------------------------------------------------------------------
h1('11. Trade-offs')

h2('11.1 Advantages')
bullet('Separation of concerns — each agent is independently testable and debuggable')
bullet('Model cost efficiency — cheap models handle cheap tasks')
bullet('County scalability — new counties require minimal changes')
bullet('Loopback is natural — Case Manager handles it with an IF node')
bullet('Context window efficiency — Title Attorney sees clean summaries, not raw HTML')
bullet('Parallelism is possible at the sub-agent dispatch level')

doc.add_paragraph()
h2('11.2 Trade-offs')
bullet('More workflows to build and maintain — 7 sub-agents vs 4 phases')
bullet('HTTP round-trip latency — each sub-agent call adds overhead vs direct tool calls')
bullet('Deeper call stack — debugging a failure requires tracing through multiple workflows')
bullet('Title Attorney still runs long — the async job pattern is still required regardless of architecture')
bullet('Initial build effort — rewriting from the current linear pipeline takes time')

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 12. Build Order
# ---------------------------------------------------------------------------
h1('12. Recommended Build Order')

steps_build = [
    ('1', 'Jobs table migration', 'Add the jobs table to scraper DB for async status tracking'),
    ('2', 'Case Manager entry point', 'Webhook that creates job, responds with job_id, triggers async chain'),
    ('3', 'Property Researcher', 'Replaces Scout — same logic, cleaner agent boundary'),
    ('4', 'Deeds Expert', 'Extracted from current Investigator — search + pull only'),
    ('5', 'Court Expert', 'Extracted from current Investigator — court search + pull only'),
    ('6', 'Document Analyst', 'Extracted from current Investigator — OCR + extraction only'),
    ('7', 'Title Attorney', 'Rebuilt Investigator — delegates retrieval, focuses on chain reasoning'),
    ('8', 'Conclusion Writer', 'Existing Conclude prompt, unchanged'),
    ('9', 'Senior Partner', 'Existing Verify prompt, unchanged'),
    ('10', 'Case Manager loopback', 'Wire Verifier verdict → IF node → re-run Title Attorney if objection_raised'),
    ('11', 'Job status endpoint', 'GET /job/{id} polling endpoint'),
    ('12', 'Parallel dispatch', 'Add parallel Deeds + Court calls where applicable'),
]

tbl5 = doc.add_table(rows=1, cols=3)
tbl5.style = 'Table Grid'
hb = tbl5.rows[0]
hb.cells[0].text = 'Step'
hb.cells[1].text = 'Component'
hb.cells[2].text = 'Notes'
for cell in hb.cells:
    cell.paragraphs[0].runs[0].bold = True
shade_row(hb, 'D9E1F2')

for step_data in steps_build:
    row = tbl5.add_row()
    for i, val in enumerate(step_data):
        row.cells[i].text = val

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 13. What Stays the Same
# ---------------------------------------------------------------------------
h1('13. What Stays the Same')

body('The following are preserved exactly from the current architecture:')
doc.add_paragraph()

bullet('Database schema — all existing tables, no migration needed except the jobs table')
bullet('All server endpoints (/investigate/*, /conclude/*, /verify/*, /scout/write, etc.)')
bullet('Document extraction subroutine and document_read logic')
bullet('Investigate Agent system prompt — carried into Title Attorney with minimal changes')
bullet('Conclude prompt — carried into Conclusion Writer unchanged')
bullet('Verify prompt — carried into Senior Partner unchanged')
bullet('All county scraper adapters (Wake, Mecklenburg, Buncombe, New Hanover)')
bullet('Business logic and NC title law rules')

doc.add_paragraph()

# ---------------------------------------------------------------------------
# Footer note
# ---------------------------------------------------------------------------
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Chain-of-Title Scraper — New Architecture Plan')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('Internal planning document — May 2026').font.size = Pt(9)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
out_path = os.path.join(os.path.dirname(__file__), 'New Architecture Plan.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
